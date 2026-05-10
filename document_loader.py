"""
Document loader — extracts text and embedded images from PDF, DOCX, or image files.
Returns a unified dict for the grading pipeline to consume.
"""

import base64
import io
from pathlib import Path

from PIL import Image

MAX_IMAGES = 10        # Claude limit is 20; leave headroom
MAX_IMAGE_PX = 1568    # Anthropic recommended max dimension


def _resize_if_needed(img: Image.Image) -> Image.Image:
    """Downscale if either dimension exceeds the recommended max."""
    w, h = img.size
    if max(w, h) > MAX_IMAGE_PX:
        scale = MAX_IMAGE_PX / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    return img


def _pil_to_b64(img: Image.Image) -> str:
    img = _resize_if_needed(img.convert("RGB"))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()


def _load_pdf(path: str) -> dict:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    text_parts, images = [], []

    for page_num, page in enumerate(doc):
        page_text = page.get_text().strip()
        if page_text:
            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

        if len(images) >= MAX_IMAGES:
            continue

        for img_info in page.get_images(full=True):
            if len(images) >= MAX_IMAGES:
                break
            xref = img_info[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:          # CMYK → RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                # Skip tiny images (icons, bullets, etc.)
                if img.width < 80 or img.height < 80:
                    continue
                images.append(_pil_to_b64(img))
            except Exception:
                continue

    page_count = len(doc)
    doc.close()
    return {
        "text": "\n\n".join(text_parts),
        "images": images,
        "image_count": len(images),
        "page_count": page_count,
        "format": "pdf",
    }


def _load_docx(path: str) -> dict:
    import docx as python_docx

    doc = python_docx.Document(path)

    # Extract paragraph text (preserving table cell text too)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    # Extract embedded images from document relationships
    images = []
    try:
        for rel in doc.part.rels.values():
            if len(images) >= MAX_IMAGES:
                break
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(img_bytes))
                    if img.width < 80 or img.height < 80:
                        continue
                    images.append(_pil_to_b64(img))
                except Exception:
                    continue
    except Exception:
        pass

    return {
        "text": "\n".join(text_parts),
        "images": images,
        "image_count": len(images),
        "page_count": None,
        "format": "docx",
    }


def _load_image(path: str) -> dict:
    img = Image.open(path)
    return {
        "text": "",
        "images": [_pil_to_b64(img)],
        "image_count": 1,
        "page_count": 1,
        "format": Path(path).suffix.lower().lstrip("."),
    }


def load(file_path: str) -> dict:
    """
    Loads a document and returns:
      text        — full extracted text
      images      — list of base64 PNG strings (≤MAX_IMAGES)
      image_count — number of images extracted
      page_count  — number of pages (None for DOCX/images)
      format      — 'pdf', 'docx', 'png', 'jpg', etc.
    """
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _load_docx(file_path)
    elif suffix in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        return _load_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: PDF, DOCX, PNG, JPG")
